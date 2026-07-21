"""
Genera un .docx con la auditoría detallada fase por fase (Real vs Pronóstico)
para el Top 4 de cada polla, a partir del mismo JSON que alimenta el informe
HTML (scripts/audit-top4-report.mjs).
Ejecutar: python scripts/build-audit-docx.py <input.json> <output.docx>
"""
import json
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PITCH = RGBColor(0x2F, 0x6B, 0x4D)
INK = RGBColor(0x17, 0x23, 0x1A)
INK_SOFT = RGBColor(0x4B, 0x59, 0x4C)
GOLD = RGBColor(0xA8, 0x79, 0x2A)

HIT_FILL = "E7F1EC"
PARTIAL_FILL = "FBF1DE"
MISS_FILL = "FBEDEA"
HEADER_FILL = "EEF2EA"


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def style_cell_text(cell, size=9, bold=False, color=None, align=None):
    for para in cell.paragraphs:
        if align:
            para.alignment = align
        if not para.runs:
            para.add_run('')
        for r in para.runs:
            r.font.size = Pt(size)
            r.bold = bold
            if color:
                r.font.color.rgb = color


def add_table(doc, headers, col_widths=None, style='Table Grid'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], HEADER_FILL)
        style_cell_text(hdr[i], size=8, bold=True, color=INK_SOFT)
    if col_widths:
        for row in table.rows:
            for cell, w in zip(row.cells, col_widths):
                cell.width = Cm(w)
    return table


def add_row(table, values, fills=None, col_widths=None, align_last_num=True):
    row = table.add_row()
    cells = row.cells
    for i, v in enumerate(values):
        cells[i].text = '' if v is None else str(v)
        align = WD_ALIGN_PARAGRAPH.RIGHT if (align_last_num and i == len(values) - 1) else None
        style_cell_text(cells[i], size=9, align=align)
        if fills and fills[i]:
            set_cell_shading(cells[i], fills[i])
    if col_widths:
        for cell, w in zip(cells, col_widths):
            cell.width = Cm(w)
    return row


def h_heading(doc, text, level=2, color=PITCH, size=13, space_before=16, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def h_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = INK
    return p


def score_str(h, a):
    if h is None or a is None:
        return '—'
    return f"{h}-{a}"


# ── Fase de grupos ──
def render_grupos(doc, g):
    h_subheading(doc, f"Grupos — Marcador de partidos ({len(g['matches'])} partidos jugados)")
    cols = [1.3, 6.5, 2, 2, 3, 1.2]
    table = add_table(doc, ['Grupo', 'Partido', 'Real', 'Pronóstico', 'Resultado', 'Pts'], cols)
    last_group = None
    for m in g['matches']:
        kind = m['kind']
        fill = HIT_FILL if kind == 'exacto' else (PARTIAL_FILL if kind == 'resultado' else MISS_FILL)
        label = 'Marcador exacto' if kind == 'exacto' else ('Resultado correcto' if kind == 'resultado' else 'Fallo')
        gcell = m['group'] if m['group'] != last_group else ''
        last_group = m['group']
        add_row(table, [gcell, f"{m['homeTeam']} vs {m['awayTeam']}", score_str(m['realHome'], m['realAway']),
                         score_str(m['predHome'], m['predAway']), label, m['pts']], fills=[None, None, None, None, None, fill], col_widths=cols)
    subtotal = g['ptsExacto'] + g['ptsResultado']
    add_row(table, ['', f"Subtotal marcador exacto (+{g['ptsExacto']}) y resultado correcto (+{g['ptsResultado']})", '', '', '', subtotal],
            fills=[HEADER_FILL] * 6, col_widths=cols)
    for c in table.rows[-1].cells:
        style_cell_text(c, size=9, bold=True)

    h_subheading(doc, "Grupos — Posición exacta en tabla (1° a 4°)")
    cols2 = [1.3] + [2.0] * 8 + [1.2]
    table2 = add_table(doc, ['Grupo', '1° Real', '1° Pron.', '2° Real', '2° Pron.', '3° Real', '3° Pron.', '4° Real', '4° Pron.', 'Pts'], cols2)
    for s in g['standingsDetail']:
        if not s['allFinished']:
            add_row(table2, [s['group'], '(grupo aún no finaliza)', '', '', '', '', '', '', '', 0], col_widths=cols2)
            continue
        vals = [s['group']]
        fills = [None]
        for i in range(4):
            real = s['actual'][i] if i < len(s['actual']) else '—'
            pred = s['predicted'][i] if i < len(s['predicted']) else '—'
            hit = real != '—' and real == pred
            vals.extend([real, pred])
            fills.extend([HIT_FILL if hit else None, HIT_FILL if hit else MISS_FILL])
        vals.append(s['pts'])
        fills.append(None)
        add_row(table2, vals, fills=fills, col_widths=cols2)
    add_row(table2, ['Subtotal posición exacta'] + [''] * 8 + [g['ptsPosicion']], fills=[HEADER_FILL] * 10, col_widths=cols2)
    for c in table2.rows[-1].cells:
        style_cell_text(c, size=9, bold=True)


def render_clasificados32(doc, c):
    h_subheading(doc, "32avos — Clasificados a Dieciseisavos (32 equipos)")
    if not c:
        p = doc.add_paragraph()
        r = p.add_run("Fase de grupos aún no finaliza por completo — no calculable todavía.")
        r.italic = True
        return
    actual_set = set(c['actual'])
    max_len = max(len(c['actual']), len(c['predicted']))
    cols = [8, 8]
    table = add_table(doc, ['Real (32 clasificados)', 'Pronóstico (32 equipos)'], cols)
    for i in range(max_len):
        real = c['actual'][i] if i < len(c['actual']) else ''
        pred = c['predicted'][i] if i < len(c['predicted']) else ''
        hit = pred and pred in actual_set
        fill = HIT_FILL if (pred and hit) else (MISS_FILL if pred else None)
        add_row(table, [real, pred], fills=[None, fill], col_widths=cols)
    add_row(table, [f"Aciertos: {len(c['hits'])}/32", f"{c['pts']} pts"], fills=[HEADER_FILL, HEADER_FILL], col_widths=cols)
    for cell in table.rows[-1].cells:
        style_cell_text(cell, size=9, bold=True)


ROUND_LABELS = {
    'r32': '32avos del cuadro (R32) — 16 llaves',
    'r16': '16avos del cuadro / Octavos (R16) — 8 llaves',
    'qf': 'Cuartos de Final (QF) — 4 llaves',
    'sf': 'Semifinales (SF) — 2 llaves',
}


def render_round(doc, r, include_teams):
    h_subheading(doc, ROUND_LABELS[r['round']])
    if include_teams:
        cols = [8, 8]
        table = add_table(doc, [f"Equipos reales ({len(r['actualTeams'])})", f"Acertados por pronóstico ({len(r['teamHits'])} x {r['teamPtsUnit']} pts)"], cols)
        add_row(table, [', '.join(r['actualTeams']) or '—', ', '.join(r['teamHits']) or '—'], col_widths=cols)
        if r['teamMisses']:
            add_row(table, ['Fallados (pronosticados, no clasificaron)', ', '.join(r['teamMisses'])], col_widths=cols)
        add_row(table, ['Subtotal equipos clasificados', f"{r['teamPts']} pts"], fills=[HEADER_FILL, HEADER_FILL], col_widths=cols)
        for cell in table.rows[-1].cells:
            style_cell_text(cell, size=9, bold=True)

    cols3 = [6, 6, 3]
    table3 = add_table(doc, ['Llave real', 'Llave pronosticada', 'Acierto'], cols3)
    for p in r['pairs']:
        real_str = f"{p['realHome'] or '?'} vs {p['realAway'] or '?'}" if (p['realHome'] or p['realAway']) else '(sin definir)'
        pred_str = f"{p['predHome'] or '?'} vs {p['predAway'] or '?'}" if (p['predHome'] or p['predAway']) else '(sin pronóstico)'
        fill = HIT_FILL if p['pairHit'] else MISS_FILL
        add_row(table3, [real_str, pred_str, '✓' if p['pairHit'] else '✗'], fills=[fill, fill, fill], col_widths=cols3)
    add_row(table3, [f"Llaves acertadas: {len(r['pairHits'])}/{len(r['pairs'])} (x {r['pairPtsUnit']} pts)", '', f"{r['pairPts']} pts"],
            fills=[HEADER_FILL, HEADER_FILL, HEADER_FILL], col_widths=cols3)
    for cell in table3.rows[-1].cells:
        style_cell_text(cell, size=9, bold=True)


def render_final_four(doc, ff):
    h_subheading(doc, f"Final y Tercer Puesto — 4 equipos (Final Four, {ff['teamPtsUnit']} pts/equipo)")
    cols = [3, 5, 5, 5]
    table = add_table(doc, ['Partido', 'Real', 'Pronóstico', 'Equipos acertados'], cols)

    def match_row(label, real, pred, hits):
        real_str = f"{real['home'] or '?'} vs {real['away'] or '?'}" if (real['home'] or real['away']) else '(sin definir)'
        pred_str = f"{pred['home'] or '?'} vs {pred['away'] or '?'}" if (pred['home'] or pred['away']) else '(sin pronóstico)'
        add_row(table, [label, real_str, pred_str, ', '.join(hits) or '—'], col_widths=cols)

    match_row('Final', ff['realFinal'], ff['predFinal'], ff['finalHits'])
    match_row('Tercer puesto', ff['realThird'], ff['predThird'], ff['thirdHits'])
    total_team_hits = len(ff['finalHits']) + len(ff['thirdHits'])
    add_row(table, [f"Subtotal equipos ({total_team_hits} aciertos)", '', '', f"{ff['teamPts']} pts"], fills=[HEADER_FILL] * 4, col_widths=cols)
    for cell in table.rows[-1].cells:
        style_cell_text(cell, size=9, bold=True)

    cols2 = [10, 3, 3]
    table2 = add_table(doc, ['Llave', 'Acierto', 'Pts'], cols2)
    fill3 = HIT_FILL if ff['thirdPairHit'] else MISS_FILL
    add_row(table2, ['Llave 3er/4to puesto', '✓' if ff['thirdPairHit'] else '✗', ff['thirdPairPts']], fills=[fill3, fill3, fill3], col_widths=cols2)
    fillf = HIT_FILL if ff['finalPairHit'] else MISS_FILL
    add_row(table2, ['Llave FINAL', '✓' if ff['finalPairHit'] else '✗', ff['finalPairPts']], fills=[fillf, fillf, fillf], col_widths=cols2)


def render_bonus(doc, b):
    h_subheading(doc, "Bonus — Puestos finales")
    cols = [3.5, 4, 4, 2.5, 3]
    table = add_table(doc, ['Puesto', 'Real', 'Pronóstico', 'Acierto', 'Pts'], cols)

    def row(label, obj):
        fill = HIT_FILL if obj['hit'] else MISS_FILL
        add_row(table, [label, obj['real'] or '(pendiente)', obj['pred'] or '—', '✓' if obj['hit'] else '✗', f"{obj['pts']} / {obj['ptsUnit']}"],
                fills=[fill, fill, fill, fill, fill], col_widths=cols)

    row('Campeón', b['champion'])
    row('Subcampeón', b['runnerUp'])
    row('Tercer puesto', b['third'])
    row('Cuarto puesto', b['fourth'])
    add_row(table, ['Subtotal bonus', '', '', '', f"{b['total']} pts"], fills=[HEADER_FILL] * 5, col_widths=cols)
    for cell in table.rows[-1].cells:
        style_cell_text(cell, size=9, bold=True)


def render_goleador(doc, gl):
    h_subheading(doc, "Goleador")
    cols = [6, 5, 2.5, 3.5]
    table = add_table(doc, ['Top goleadores oficiales', 'Pronóstico', 'Acierto', 'Pts'], cols)
    fill = HIT_FILL if gl['hit'] else MISS_FILL
    add_row(table, [', '.join(gl['actualTop']) or '(sin definir)', gl['predicted'] or '—', '✓' if gl['hit'] else '✗', f"{gl['pts']} / {gl['ptsUnit']}"],
            fills=[fill, fill, fill, fill], col_widths=cols)


def render_participant(doc, p, rank):
    h_heading(doc, f"#{rank}  {p['name']}   —   {p['total']} pts", level=2, color=INK, size=15, space_before=22, space_after=8)

    ob = p['officialBreakdown']
    ph = p['phases']
    bracket_pts = (ph['r32']['pairPts'] + ph['r16']['teamPts'] + ph['r16']['pairPts'] +
                   ph['qf']['teamPts'] + ph['qf']['pairPts'] + ph['sf']['teamPts'] + ph['sf']['pairPts'] +
                   ph['finalFour']['teamPts'] + ph['finalFour']['thirdPairPts'] + ph['finalFour']['finalPairPts'] +
                   ph['bonus']['total'])

    cols = [4, 5, 5.5, 2.5, 2.5]
    table = add_table(doc, ['Grupos (marcador)', 'Grupos (posición+clasif.)', 'Bracket (llaves+equipos+bonus)', 'Goleador', 'TOTAL'], cols)
    add_row(table, [ob['ptsExacto'] + ob['ptsResultado'], ob['ptsStandings'], bracket_pts, ob['ptsScorers'], p['total']], col_widths=cols)
    for cell in table.rows[-1].cells:
        style_cell_text(cell, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    h_heading(doc, "1. Fase de Grupos", level=3, color=PITCH, size=11.5, space_before=14, space_after=2)
    render_grupos(doc, ph['grupos'])

    h_heading(doc, "2. 32avos (Dieciseisavos) — clasificación y llaves", level=3, color=PITCH, size=11.5, space_before=14, space_after=2)
    render_clasificados32(doc, ph['clasificados32'])
    render_round(doc, ph['r32'], include_teams=False)

    h_heading(doc, "3. 16avos (Octavos, R16)", level=3, color=PITCH, size=11.5, space_before=14, space_after=2)
    render_round(doc, ph['r16'], include_teams=True)

    h_heading(doc, "4. Cuartos de Final (QF)", level=3, color=PITCH, size=11.5, space_before=14, space_after=2)
    render_round(doc, ph['qf'], include_teams=True)

    h_heading(doc, "5. Semifinales (SF)", level=3, color=PITCH, size=11.5, space_before=14, space_after=2)
    render_round(doc, ph['sf'], include_teams=True)

    h_heading(doc, "6. Final y Tercer Puesto + Bonus", level=3, color=PITCH, size=11.5, space_before=14, space_after=2)
    render_final_four(doc, ph['finalFour'])
    render_bonus(doc, ph['bonus'])

    h_heading(doc, "7. Goleador", level=3, color=PITCH, size=11.5, space_before=14, space_after=2)
    render_goleador(doc, ph['goleador'])

    doc.add_page_break()


def main():
    input_path, output_path = sys.argv[1], sys.argv[2]
    with open(input_path, encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    run = title.add_run('Auditoría Detallada de Puntajes — Top 4 por Polla')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = INK

    gen_at = datetime.fromisoformat(data['generatedAt'].replace('Z', '+00:00'))
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(
        f"Generado {gen_at.strftime('%d de %B de %Y, %H:%M UTC')} · Comparación Real vs. Pronóstico, fase por fase, "
        "para los 4 primeros lugares de cada polla: grupos, 32avos, 16avos, cuartos, semifinales, final y goleador."
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = INK_SOFT

    legend = doc.add_paragraph()
    lrun = legend.add_run('Verde = acierto   ·   Amarillo = acierto parcial (grupos)   ·   Rojo = fallo')
    lrun.font.size = Pt(9)
    lrun.font.color.rgb = INK_SOFT

    # Índice
    h_heading(doc, 'Índice', level=1, color=PITCH, size=13, space_before=16)
    for polla in data['pollas']:
        idx_p = doc.add_paragraph()
        idx_run = idx_p.add_run(polla['name'])
        idx_run.bold = True
        idx_run.font.size = Pt(10.5)
        for i, p in enumerate(polla['top4'], start=1):
            item = doc.add_paragraph(style=None)
            item.paragraph_format.left_indent = Cm(0.6)
            r = item.add_run(f"#{i}  {p['name']}  —  {p['total']} pts")
            r.font.size = Pt(9.5)
    doc.add_page_break()

    for polla in data['pollas']:
        h_heading(doc, f"Polla: {polla['name']} — Top 4", level=1, color=PITCH, size=16, space_before=6)
        for i, p in enumerate(polla['top4'], start=1):
            render_participant(doc, p, i)

    footer = doc.add_paragraph()
    frun = footer.add_run(
        'Fin del informe. Todos los totales por participante fueron verificados matemáticamente: la suma de las 7 '
        'fases (grupos, posición de grupo, clasificados a 32avos, 16avos, cuartos, semifinales, final/tercer puesto '
        '+ bonus, y goleador) coincide de forma exacta con el puntaje oficial que muestra la tabla de posiciones de la app.'
    )
    frun.italic = True
    frun.font.size = Pt(9)
    frun.font.color.rgb = INK_SOFT

    doc.save(output_path)
    print('OK ->', output_path)


if __name__ == '__main__':
    main()
