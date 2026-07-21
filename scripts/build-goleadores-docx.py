"""
Genera un .docx con el listado Participante / Goleador pronosticado a partir
de scratchpad/goleadores.json (mismos datos que el artifact HTML publicado).
Ejecutar: python scripts/build-goleadores-docx.py <input.json> <output.docx>
"""
import json
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

STATUS_LABEL = {
    'corregido': 'Corregido 2026-07-20',
    'ya-exacto': 'Ya era exacto',
    'exacto-otro': 'Coincide (otro jugador)',
    'sin-cambio': 'Sin cambios',
    'sin-pronostico': 'Sin pronóstico',
}

PITCH = RGBColor(0x2F, 0x6B, 0x4D)
BRICK = RGBColor(0xA8, 0x46, 0x32)
INK_SOFT = RGBColor(0x4B, 0x59, 0x4C)
GOLD = RGBColor(0xA8, 0x79, 0x2A)

HIT_FILL = "E7F1EC"
MISS_FILL = "FBEDEA"


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def add_header_row(table, headers):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = INK_SOFT
        set_cell_shading(hdr[i], "EEF2EA")


def main():
    input_path, output_path = sys.argv[1], sys.argv[2]
    with open(input_path, encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()

    # Márgenes y estilo base
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)

    # Título
    title = doc.add_paragraph()
    run = title.add_run('Listado Goleador Pronosticado — Validación')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x17, 0x23, 0x1A)

    gen_at = datetime.fromisoformat(data['generatedAt'].replace('Z', '+00:00'))
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(
        f"Generado {gen_at.strftime('%d de %B de %Y, %H:%M UTC')} · "
        f"Goleador real: {', '.join(data['topScorers'])} · "
        "Participante / Goleador pronosticado por cada uno, validación de la normalización aplicada el 2026-07-20."
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = INK_SOFT

    # Resumen
    total = len(data['list'])
    hits = sum(1 for p in data['list'] if p.get('hit'))
    fixed = sum(1 for p in data['list'] if p.get('status') == 'corregido')

    summary = doc.add_paragraph()
    summary_run = summary.add_run(
        f"Total participantes: {total}    |    Acertaron goleador (+10 pts): {hits}    |    Corregidos el 2026-07-20: {fixed}"
    )
    summary_run.bold = True
    summary_run.font.size = Pt(11)

    legend = doc.add_paragraph()
    legend_run = legend.add_run(
        'Fila verde = acertó (cuenta +10 pts) · Fila roja = no acertó · "(Corregido)" junto al nombre = su registro se '
        'normalizó a "Kylian Mbappé" en la sesión del 2026-07-20.'
    )
    legend_run.italic = True
    legend_run.font.size = Pt(9)
    legend_run.font.color.rgb = INK_SOFT

    pollas = []
    for p in data['list']:
        if p['polla'] not in pollas:
            pollas.append(p['polla'])

    for polla in pollas:
        rows = [p for p in data['list'] if p['polla'] == polla]
        polla_hits = sum(1 for p in rows if p.get('hit'))

        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        hrun = h.add_run(f"{polla}  ")
        hrun.bold = True
        hrun.font.size = Pt(14)
        hrun.font.color.rgb = PITCH
        hrun2 = h.add_run(f"({len(rows)} participantes · {polla_hits} con goleador acertado)")
        hrun2.font.size = Pt(9)
        hrun2.font.color.rgb = INK_SOFT

        table = doc.add_table(rows=1 + len(rows), cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        add_header_row(table, ['Participante', 'Goleador pronosticado', 'Acierto'])
        set_col_widths(table, [6.5, 6.5, 2.5])

        for i, p in enumerate(rows, start=1):
            cells = table.rows[i].cells
            name = p['name']
            if p.get('status') == 'corregido':
                name += '  (Corregido)'
            cells[0].text = name
            goleador = p.get('goleador') or '(sin pronóstico)'
            cells[1].text = goleador
            cells[2].text = '✓' if p.get('hit') else ('—' if p.get('status') == 'sin-pronostico' else '✗')
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for c in cells:
                for para in c.paragraphs:
                    for r in para.runs:
                        r.font.size = Pt(10)
                    if not para.runs:
                        para.font = normal.font

            if p.get('hit'):
                for c in cells:
                    set_cell_shading(c, HIT_FILL)
            elif p.get('status') != 'sin-pronostico':
                for c in cells:
                    set_cell_shading(c, MISS_FILL)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(20)
    frun = footer.add_run(
        'Fin del listado. Los 41 aciertos de "Kylian Mbappé" ya están reflejados en la base de datos '
        '(11 que ya venían exactos + 30 normalizados el 2026-07-20). Cualquier otro nombre no confirmado como '
        'variante válida (ej. "kanw") quedó sin tocar.'
    )
    frun.italic = True
    frun.font.size = Pt(9)
    frun.font.color.rgb = INK_SOFT

    doc.save(output_path)
    print('OK ->', output_path)


if __name__ == '__main__':
    main()
